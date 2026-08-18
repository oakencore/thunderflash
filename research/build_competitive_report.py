from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(__file__).with_name("thunderflash_competitive_landscape.docx")

NAVY = "17324D"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
DARK = "1F2937"
GREEN = "1F6B52"
AMBER = "8A6116"
RED = "9B1C1C"
WHITE = "FFFFFF"


def set_font(run, name="Calibri", size=11, bold=None, color=DARK, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[i]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[i] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def table_borders(table, color="D0D5DD", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)


def set_cell_text(cell, text, size=9, bold=False, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    set_font(p.add_run(text), size=size, bold=bold, color=color)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.167
    set_font(p.add_run(text), size=11)
    return p


def add_link(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_source(doc, n, label, title, url):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    set_font(p.add_run(f"[{n}] {label}. "), size=9, bold=True, color=NAVY)
    if url.startswith("internal:"):
        set_font(p.add_run(title), size=9, color=DARK)
    else:
        add_link(p, title, url)


def add_callout(doc, label, text, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    table_borders(table, color=accent, size="10")
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(label + "  "), size=10.5, bold=True, color=accent)
    set_font(p.add_run(text), size=10.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

# standard_business_brief token map
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for style_name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
]:
    st = styles[style_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

if "Table Citation" not in [s.name for s in styles]:
    citation = styles.add_style("Table Citation", WD_STYLE_TYPE.PARAGRAPH)
else:
    citation = styles["Table Citation"]
citation.font.name = "Calibri"
citation.font.size = Pt(8.5)
citation.font.color.rgb = RGBColor.from_string(MID_GRAY)
citation.paragraph_format.space_before = Pt(4)
citation.paragraph_format.space_after = Pt(4)

# Running furniture
hp = sec.header.paragraphs[0]
hp.text = "THUNDERFLASH  /  COMPETITIVE LANDSCAPE"
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(hp.runs[0], size=8.5, bold=True, color=MID_GRAY)
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(fp.add_run("Research snapshot  |  18 August 2026"), size=8, color=MID_GRAY)

# Editorial cover pattern
doc.add_paragraph().paragraph_format.space_after = Pt(44)
k = doc.add_paragraph()
k.alignment = WD_ALIGN_PARAGRAPH.CENTER
k.paragraph_format.space_after = Pt(14)
set_font(k.add_run("COMPETITIVE RESEARCH"), size=10, bold=True, color=AMBER)
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_after = Pt(8)
set_font(t.add_run("Thunderflash"), size=30, bold=True, color=NAVY)
st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
st.paragraph_format.space_after = Pt(6)
set_font(st.add_run("Feature Review & Competitive Landscape"), size=17, color=BLUE)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(30)
set_font(sub.add_run("How a purpose-built Thunderbolt transfer tool compares with the products people already reach for"), size=11.5, italic=True, color=MID_GRAY)

meta = doc.add_table(rows=3, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_geometry(meta, [2200, 5200], indent=0)
for i, (label, value) in enumerate([
    ("Prepared for", "Thunderflash product team"),
    ("Research date", "18 August 2026"),
    ("Decision lens", "One-off, very large Mac-to-Mac transfers over a Thunderbolt cable"),
]):
    set_cell_text(meta.cell(i, 0), label.upper(), size=8.5, bold=True, color=MID_GRAY)
    set_cell_text(meta.cell(i, 1), value, size=10.5, color=DARK)
    if i % 2 == 0:
        shade(meta.cell(i, 0), LIGHT_GRAY)
        shade(meta.cell(i, 1), LIGHT_GRAY)
table_borders(meta, color="D0D5DD", size="4")

doc.add_page_break()

doc.add_heading("Executive summary", level=1)
lead = doc.add_paragraph()
lead.paragraph_format.space_after = Pt(10)
set_font(lead.add_run("Bottom line. "), size=11.5, bold=True, color=NAVY)
set_font(lead.add_run("Thunderflash has a credible wedge: it removes the setup and protocol overhead that keeps general-purpose tools from feeling native to a direct Thunderbolt cable. For that exact job, it is simpler than rsync and more purpose-built than Finder, AirDrop, or internet-oriented handoff tools."), size=11.5)

add_callout(doc, "POSITION TODAY", "Best-in-class focus and likely link utilization for trusted, cable-only, one-time bulk transfer; incomplete as a dependable general-purpose transfer product.", fill="EAF4F0", accent=GREEN)

add_bullet(doc, "Where Thunderflash leads: cable-specific discovery, one pipelined TCP stream, no SSH encryption overhead, live throughput, automatic keep-awake behavior, a fresh three-word pairing phrase, and no background service or account.")
add_bullet(doc, "Where it trails: no transport encryption, resume, end-to-end content hash, xattrs/ACLs/resource forks, pull mode, GUI, cross-platform support, or remote/LAN use beyond the Thunderbolt bridge.")
add_bullet(doc, "The strongest direct substitutes are Finder File Sharing over Thunderbolt Bridge, Target Disk Mode/Share Disk, and rsync over the bridge. The strongest habit competitors are AirDrop and Migration Assistant. The closest phrase-driven CLI alternative is croc.")
add_bullet(doc, "The highest-leverage roadmap is integrity verification first, then resume, then macOS metadata fidelity. Encryption is important for scope expansion, but is less urgent if the product remains explicitly cable-only and trusted-device-only.")

doc.add_page_break()
doc.add_heading("What Thunderflash does today", level=1)
features = [
    ("Purpose-built transport", "Uses the macOS Thunderbolt Bridge interface directly and refuses to fall back to another interface."),
    ("Low-friction pairing", "Receiver advertises its address and a fresh, single-use three-word phrase; sender can discover it or connect directly."),
    ("Bulk-transfer performance", "Streams files and directories over one pipelined TCP connection with large chunks and avoids per-file round trips."),
    ("Operational feedback", "Shows file counts, bytes, instantaneous throughput, average throughput, current path, and completion statistics."),
    ("Mac-aware ergonomics", "Uses caffeinate for the lifetime of a transfer and warns when bridge MTU remains at 1500."),
    ("Filesystem basics", "Transfers files, directories, empty entries, and symlinks; preserves Unix mode and modification time."),
    ("Receiver hardening", "Rejects absolute paths, parent traversal and NULs; uses no-follow opens so transmitted or pre-existing symlinks cannot redirect writes outside the destination."),
    ("Deliberate limits", "Push only; one sender per receive run; no daemon; no encryption, resume, content hashing, xattrs, ACLs, or resource forks."),
]
ft = doc.add_table(rows=1, cols=2)
ft.alignment = WD_TABLE_ALIGNMENT.LEFT
headers = ["Capability", "Current behavior"]
for i, h in enumerate(headers):
    set_cell_text(ft.cell(0, i), h, size=9, bold=True, color=WHITE)
    shade(ft.cell(0, i), NAVY)
for label, desc in features:
    row = ft.add_row()
    set_cell_text(row.cells[0], label, size=9.3, bold=True, color=NAVY)
    set_cell_text(row.cells[1], desc, size=9.3)
set_repeat_table_header(ft.rows[0])
set_table_geometry(ft, [2300, 7060])
table_borders(ft)
p = doc.add_paragraph(style="Table Citation")
p.add_run("Source: Thunderflash README, CLI implementation, transfer protocol, and repository tests reviewed 18 August 2026.")

doc.add_page_break()
doc.add_heading("The tools users are most likely to choose", level=1)
p = doc.add_paragraph("The market is best understood in three groups. This avoids penalizing a convenient handoff tool for not being a cable optimizer—or giving a complex sync platform credit for a workflow it makes cumbersome.")

market = [
    ("Thunderflash", "Direct specialist", "Purpose-built Thunderbolt push", "Very low", "Trusted, very large one-off transfer", "Missing resilience and full-fidelity safeguards"),
    ("Finder File Sharing", "Direct substitute", "SMB over Thunderbolt Bridge", "Med.", "Native browsing and copy without extra software", "Sharing/login setup; less focused CLI workflow"),
    ("Target Disk Mode / Share Disk", "Direct substitute", "Mac exposes storage to the other Mac", "Med.", "Native bulk access, migration, or recovery", "Requires restart/recovery workflow; not a lightweight live handoff"),
    ("rsync", "Direct substitute", "Remote shell or daemon over bridge IP", "High", "Technical users needing resume, verification, filters, fidelity", "More setup and flags; SSH may reduce peak throughput"),
    ("Migration Assistant", "Workflow substitute", "Apple migration workflow", "Low-med.", "Moving accounts, apps, settings, and data to a new Mac", "Overkill for ad hoc selected files"),
    ("AirDrop", "Habit competitor", "Nearby wireless Apple transfer", "Very low", "Small-to-medium casual sends", "Not designed for sustained multi-terabyte wired transfers"),
    ("croc", "Experience competitor", "Direct or relay; code phrase", "Low", "Cross-platform, encrypted, resumable one-off transfer", "General-purpose routing and crypto add overhead; not Thunderbolt-specific"),
    ("Magic Wormhole", "Experience competitor", "Direct TCP or encrypted relay; short code", "Low-med.", "Secure ad hoc transfer across networks", "Relay/rendezvous dependency and less cable-specific optimization"),
    ("LocalSend", "Habit competitor", "Local-network GUI transfer", "Low", "Offline, cross-platform nearby sharing", "GUI/LAN focus; speed bounded by chosen LAN path"),
    ("Taildrop", "Ecosystem substitute", "Fastest Tailscale path", "Med.", "Encrypted sends among a user's enrolled devices", "Requires Tailscale identity and device enrollment"),
    ("Syncthing", "Workflow substitute", "Continuous P2P block sync", "High initially", "Ongoing folder synchronization and versioning", "Persistent configuration is disproportionate for one-time handoff"),
]
mt = doc.add_table(rows=1, cols=6)
for i, h in enumerate(["Tool", "Role", "Primary path", "Setup", "Best fit", "Main tradeoff"]):
    set_cell_text(mt.cell(0, i), h, size=8.1, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER if i == 3 else WD_ALIGN_PARAGRAPH.LEFT)
    shade(mt.cell(0, i), NAVY)
for ridx, rowdata in enumerate(market):
    row = mt.add_row()
    for i, val in enumerate(rowdata):
        set_cell_text(row.cells[i], val, size=7.7, bold=(i == 0), color=NAVY if i == 0 else DARK, align=WD_ALIGN_PARAGRAPH.CENTER if i == 3 else WD_ALIGN_PARAGRAPH.LEFT)
        if ridx % 2 == 1:
            shade(row.cells[i], LIGHT_GRAY)
set_repeat_table_header(mt.rows[0])
set_table_geometry(mt, [1250, 1200, 1600, 700, 2220, 2390])
table_borders(mt, size="4")

doc.add_page_break()
doc.add_heading("Capability comparison", level=1)
doc.add_paragraph("Legend: Yes = built in; Partial = possible with caveats, extra flags, or a different product mode; No = not a normal capability. “Cable-first” means the workflow deliberately selects or exposes the Thunderbolt path, not merely that it could route over any IP interface.")

matrix_rows = [
    ("Thunderflash", "Yes", "No", "No", "Length only", "Mode + mtime", "Phrase", "Push"),
    ("Finder / SMB", "Yes", "Yes", "Partial", "Protocol checks", "Strong on Mac", "Account", "Both"),
    ("Target Disk / Share Disk", "Yes", "Depends", "N/A", "Filesystem", "Strong on Mac", "Physical / login", "Both"),
    ("rsync", "Via config", "Via SSH", "Yes", "Whole-file checksum", "Strong with flags", "SSH / daemon", "Both"),
    ("Migration Assistant", "Partial", "Apple-managed", "Retry", "Apple-managed", "Very strong", "Login / code", "Import"),
    ("AirDrop", "No", "Yes", "Opaque", "Managed", "Good for items", "Proximity / consent", "Push"),
    ("croc", "No", "Yes (PAKE)", "Yes", "Hash verification", "Basic", "Code phrase", "Push"),
    ("Magic Wormhole", "No", "Yes (PAKE)", "No*", "SHA-256 ack", "Basic", "Short code", "Push"),
    ("LocalSend", "No", "Yes", "Opaque", "Managed", "Basic", "Device consent", "Both"),
    ("Taildrop", "No", "Yes", "Partial", "Managed", "Basic", "Tailscale identity", "Push"),
    ("Syncthing", "No", "Yes (TLS)", "Yes", "SHA-256 blocks", "Strong sync semantics", "Mutual device IDs", "Both"),
]
cm = doc.add_table(rows=1, cols=8)
for i, h in enumerate(["Tool", "Cable-first", "Encrypted", "Resume", "Integrity", "Metadata", "Pairing", "Flow"]):
    set_cell_text(cm.cell(0, i), h, size=7.7, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade(cm.cell(0, i), BLUE)
for ridx, vals in enumerate(matrix_rows):
    row = cm.add_row()
    for i, val in enumerate(vals):
        color = DARK
        if val == "Yes" or val.startswith("Yes ("):
            color = GREEN
        elif val == "No" or val.startswith("No/"):
            color = RED
        set_cell_text(row.cells[i], val, size=7.4, bold=(i == 0), color=NAVY if i == 0 else color, align=WD_ALIGN_PARAGRAPH.CENTER if i != 0 else WD_ALIGN_PARAGRAPH.LEFT)
        if ridx % 2 == 1:
            shade(row.cells[i], LIGHT_GRAY)
set_repeat_table_header(cm.rows[0])
set_table_geometry(cm, [1400, 900, 980, 760, 1350, 1420, 1740, 810])
table_borders(cm, size="4")
cp = doc.add_paragraph(style="Table Citation")
cp.add_run("* Magic Wormhole can reconnect its control protocol while a process stays alive, but its standard one-shot file-transfer workflow should not be treated as durable cross-process file resume. Product behaviors can vary by OS and version; validate before publishing hard claims.")

doc.add_page_break()
doc.add_heading("How Thunderflash stacks up", level=1)
stack = [
    ("Raw fit for the target job", "LEADS", "It is the only tool in this set designed around a live Thunderbolt Bridge, a single ephemeral receive session, and minimal per-file protocol work."),
    ("Time to first transfer", "LEADS / TIES", "Once installed, the two-command phrase flow is simpler than SMB sharing and rsync setup; AirDrop remains easier for casual small sends."),
    ("Peak-speed potential", "LEADS", "Plain TCP, large chunks, no SSH crypto, and no per-file handshakes are a strong architecture for saturating the chosen cable path. The repository's 1.2–2.5 GB/s figure is an expectation, not an independently verified benchmark."),
    ("Security scope", "TRAILS", "Interface binding and a one-guess phrase limit exposure, but plaintext data and a roughly 24-bit phrase are not substitutes for authenticated encryption."),
    ("Failure recovery", "TRAILS", "A dropped multi-terabyte transfer starts over. rsync, croc, and Syncthing have materially stronger restart behavior."),
    ("Data confidence", "TRAILS", "Length checks catch truncation but not silent corruption. Several alternatives verify transferred content cryptographically or by block hash."),
    ("Mac fidelity", "TRAILS", "Missing xattrs, ACLs, and resource forks makes the current build unsafe for app bundles, Photos libraries, and other macOS-rich objects."),
    ("Breadth", "TRAILS BY DESIGN", "Mac-only, cable-only, push-only, and CLI-only is a sharp wedge, but it constrains addressable use cases compared with cross-platform and remote tools."),
]
stbl = doc.add_table(rows=1, cols=3)
for i, h in enumerate(["Dimension", "Position", "Assessment"]):
    set_cell_text(stbl.cell(0, i), h, size=9, bold=True, color=WHITE)
    shade(stbl.cell(0, i), NAVY)
for ridx, (dim, pos, assess) in enumerate(stack):
    row = stbl.add_row()
    set_cell_text(row.cells[0], dim, size=8.8, bold=True, color=NAVY)
    c = GREEN if "LEAD" in pos else RED
    set_cell_text(row.cells[1], pos, size=8, bold=True, color=c, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[2], assess, size=8.8)
    if ridx % 2 == 1:
        for cell in row.cells:
            shade(cell, LIGHT_GRAY)
set_repeat_table_header(stbl.rows[0])
set_table_geometry(stbl, [1900, 1350, 6110])
table_borders(stbl)

doc.add_page_break()
doc.add_heading("Competitor notes", level=1)
notes = [
    ("Finder File Sharing over Thunderbolt Bridge", "The closest built-in substitute. Apple documents using IP over Thunderbolt, then connecting in Finder through Network after enabling File Sharing. It benefits from native UI and standard authentication. Thunderflash's advantage is a purpose-built, ephemeral transfer session with explicit throughput and no persistent share configuration."),
    ("Target Disk Mode / Share Disk", "The strongest native option when one Mac can be restarted or placed in recovery and exposed as storage. It is excellent for migration and recovery, but interrupts normal use and is less convenient for a quick live push."),
    ("rsync over the bridge", "The strongest technical competitor. It brings resume/partial-file support, transferred-file checksum verification, filters, bidirectional workflows, and extensive preservation flags. Thunderflash can still win when users want the fastest simple whole-file handoff and do not want SSH, daemon, or flag complexity."),
    ("Migration Assistant", "Owns the 'move everything to a new Mac' job: accounts, applications, settings, and files. Thunderflash should not chase this breadth; it should position around selected, extremely large payloads between Macs that remain in active use."),
    ("AirDrop", "The default mental model for nearby Apple sharing and the easiest UX benchmark. Its wireless, consent-driven workflow is ideal for everyday files. Thunderflash should borrow its confidence and simplicity while owning workloads where duration, size, and predictable wired throughput matter."),
    ("croc", "The closest experience competitor: cross-platform files/folders, a human code, PAKE-based end-to-end encryption, relay support, resume, and hash verification. Thunderflash's differentiation is local cable specialization and less overhead; croc exposes the minimum trust and resilience users may expect from a modern phrase-based transfer tool."),
    ("Magic Wormhole", "A mature short-code model with PAKE, encrypted transit, direct TCP attempts, and relay fallback. It is broader and safer across untrusted networks, but depends on rendezvous infrastructure and is not optimized specifically for Thunderbolt Bridge."),
    ("LocalSend", "A polished, offline, cross-platform LAN handoff experience with end-to-end encryption. It competes for users who prefer a GUI and do not want accounts. Thunderflash remains differentiated by explicit cable selection and a CLI optimized for huge transfers."),
    ("Tailscale Taildrop", "A convenient encrypted path between a user's already-enrolled devices, selecting the fastest available Tailscale route. It is compelling for existing Tailscale users, but account/device enrollment is heavier than Thunderflash's one-time phrase."),
    ("Syncthing", "A continuous synchronization system, not a one-shot courier. It offers encrypted transport, block hashing, changed-block transfer, mutual device identity, and optional versioning. Its persistent folder/device model is unnecessary overhead for Thunderflash's core job."),
]
for title, text in notes:
    if title == "LocalSend":
        doc.add_page_break()
    doc.add_heading(title, level=2)
    doc.add_paragraph(text)

strategy_heading = doc.add_heading("Recommended product strategy", level=1)
add_callout(doc, "PROTECT THE WEDGE", "Keep Thunderflash opinionated: two Macs, one cable, one transfer, maximum useful throughput. Add safeguards that make that promise dependable before expanding the market surface.", fill="EAF4F0", accent=GREEN)

doc.add_heading("Roadmap priority", level=2)
roadmap = [
    ("1", "End-to-end content verification", "Add a fast whole-file or chunk hash and report verified completion. This closes the largest trust gap without changing the core workflow."),
    ("2", "Resume with stable manifests", "Persist a transfer manifest and partial-file state so interrupted multi-terabyte jobs continue safely. This is the strongest competitive gap versus rsync and croc."),
    ("3", "macOS metadata fidelity", "Preserve xattrs, ACLs, flags, resource forks, sparse-file semantics, and hard links, or explicitly package rich objects. Until then, keep strong warnings around .app bundles and Photos libraries."),
    ("4", "Authenticated encryption option", "Offer a high-speed encrypted mode with integrity protection. Keep trusted-cable performance mode only if the risk is explicit and the interface boundary remains strict."),
    ("5", "Preflight and postflight confidence", "Estimate required space, detect filename conflicts, summarize skipped special files, verify receiver writes, and produce a machine-readable transfer report."),
    ("6", "Measured benchmark program", "Publish repeatable results by cable generation, Mac model, filesystem, file mix, MTU, encryption mode, and competing tool configuration. Avoid unsupported 'fastest' claims."),
]
rt = doc.add_table(rows=1, cols=3)
for i, h in enumerate(["Priority", "Capability", "Why it matters"]):
    set_cell_text(rt.cell(0, i), h, size=9, bold=True, color=WHITE)
    shade(rt.cell(0, i), NAVY)
for ridx, vals in enumerate(roadmap):
    row = rt.add_row()
    set_cell_text(row.cells[0], vals[0], size=11, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[1], vals[1], size=9, bold=True, color=NAVY)
    set_cell_text(row.cells[2], vals[2], size=9)
    if ridx % 2 == 1:
        for cell in row.cells:
            shade(cell, LIGHT_GRAY)
set_repeat_table_header(rt.rows[0])
set_table_geometry(rt, [800, 2600, 5960])
table_borders(rt)

doc.add_heading("Positioning that is supportable now", level=2)
add_bullet(doc, "A fast, temporary transfer lane for two Macs connected by Thunderbolt.")
add_bullet(doc, "No account, no daemon, no persistent share, and no SSH setup.")
add_bullet(doc, "Built for very large files and directory trees when both Macs are trusted and physically connected.")
add_bullet(doc, "Avoid claiming universal 'fastest file transfer' status until reproducible third-party-style benchmarks exist.")
add_bullet(doc, "Do not imply backup-grade or migration-grade fidelity until resume, verification, and macOS metadata support ship.")

doc.add_heading("Suggested competitive message", level=2)
q = doc.add_paragraph()
q.paragraph_format.left_indent = Inches(0.35)
q.paragraph_format.right_indent = Inches(0.35)
q.paragraph_format.space_before = Pt(8)
q.paragraph_format.space_after = Pt(8)
set_font(q.add_run("“AirDrop is for convenience. Migration Assistant is for moving a Mac. rsync is for operators. Thunderflash is the purpose-built fast lane for moving very large data between two Macs over a cable.”"), size=13, italic=True, color=NAVY)

doc.add_page_break()
doc.add_heading("Research caveats", level=1)
add_bullet(doc, "No Thunderbolt hardware benchmark was run in this environment. Performance assessments are architectural and qualitative; the 1.2–2.5 GB/s figure comes from Thunderflash's own README.")
add_bullet(doc, "Apple and third-party behavior can change with OS and product versions. The comparison reflects official documentation available on 18 August 2026.")
add_bullet(doc, "Metadata preservation labels are intentionally broad because behavior depends on source/destination filesystems, flags, transport configuration, and object type.")
add_bullet(doc, "Repository tests were reviewed and executed. Non-network tests largely passed; four loopback/network tests were blocked by the managed sandbox's socket permissions, so no product defect is inferred from those failures.")

doc.add_page_break()
doc.add_heading("Sources", level=1)
doc.add_paragraph("Primary product documentation and project-maintained references were preferred. Internal claims about Thunderflash come from the checked-out repository.")
sources = [
    ("Thunderflash", "Internal repository: README, CLI, protocol, and tests", "internal:README.md"),
    ("Apple", "Use IP over Thunderbolt to connect Mac computers", "https://support.apple.com/en-gb/guide/mac-help/mchld53dd2f5/mac"),
    ("Apple", "Transfer files using Target Disk Mode", "https://support.apple.com/en-gu/guide/mac-help/mchlp1443/mac"),
    ("Apple", "Transfer to a new Mac with Migration Assistant", "https://support.apple.com/en-gb/102613"),
    ("Apple", "Use AirDrop to send items to nearby Apple devices", "https://support.apple.com/en-za/guide/mac-help/mh35868/mac"),
    ("rsync project", "rsync(1) manual", "https://download.samba.org/pub/rsync/rsync.1"),
    ("croc project", "croc README and feature overview", "https://github.com/schollz/croc"),
    ("Magic Wormhole", "Welcome and security model", "https://magic-wormhole.readthedocs.io/en/latest/welcome.html"),
    ("Magic Wormhole", "File-transfer protocol", "https://magic-wormhole.readthedocs.io/en/latest/file-transfer-protocol.html"),
    ("LocalSend", "Official product overview", "https://redesign.localsend.org/"),
    ("Tailscale", "Taildrop documentation", "https://tailscale.com/docs/features/taildrop"),
    ("Syncthing", "Project overview", "https://github.com/syncthing"),
    ("Syncthing", "Understanding synchronization", "https://docs.syncthing.net/users/syncing"),
    ("Syncthing", "File versioning", "https://docs.syncthing.net/users/versioning"),
]
for i, (label, title, url) in enumerate(sources, 1):
    add_source(doc, i, label, title, url)

# Keep table rows together where feasible, without fixed heights.
for table in doc.tables:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)

doc.core_properties.title = "Thunderflash Feature Review & Competitive Landscape"
doc.core_properties.subject = "Competitive research for Mac-to-Mac Thunderbolt file transfer"
doc.core_properties.author = "Thunderflash product research"
doc.core_properties.keywords = "Thunderflash, Thunderbolt, file transfer, AirDrop, croc, rsync, competitive analysis"
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
